from docx import Document
import logging
import os

# Set up logging
logging.basicConfig(
    filename=r"C:\RNOS\HCP Train 2024\HCP Python\Output Folder\processing_log.txt",
    level=logging.DEBUG,  # Set to DEBUG for detailed log output
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Define the input and output file paths
input_file = r"C:\RNOS\HCP Train 2024\HCP Python\RNOS1_Abbate Rachel_Bowel and Bladder Management Healthcare Plan.docx"
output_folder = r"C:\RNOS\HCP Train 2024\HCP Python\Output Folder"

# Ensure the output folder exists
os.makedirs(output_folder, exist_ok=True)

# Construct the output file path
output_file = os.path.join(output_folder, "Updated_" + os.path.basename(input_file))

try:
    # Load the document
    doc = Document(input_file)
    logging.info(f"Document loaded: {input_file}")

    # Step 1: Modify the first row in the table, avoiding row insertions
    table = doc.tables[0]  # Assuming the first table in the document
    first_row = table.rows[0]  # Get the first row
    logging.debug(f"First row text: {[cell.text for cell in first_row.cells]}")

    # Step 2: Add or modify content within the first row instead of inserting a new row
    for idx, cell in enumerate(first_row.cells):
        new_paragraph = cell.add_paragraph("Updated content here.")
        logging.debug(f"Updated first row cell {idx} with new content.")

    # Step 3: Insert Effective Date in the corresponding row if available
    effective_date_found = False
    for row in table.rows:
        if 'Effective Date' in row.cells[0].text:
            row.cells[1].text = "10/07/2021"
            logging.info("Inserted Effective Date: 10/07/2021")
            effective_date_found = True
            break
    if not effective_date_found:
        logging.warning("Effective Date row not found in the table.")

    # Save the updated document
    doc.save(output_file)
    logging.info(f"Document saved successfully to: {output_file}")
    print("Document updated and saved.")

except Exception as e:
    logging.error(f"An error occurred: {e}")
    print(f"An error occurred: {e}")
