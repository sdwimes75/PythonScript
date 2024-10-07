import pandas as pd
from docx import Document
import openpyxl
import os
import re
import gc  # For garbage collection

# Paths
database_path = r'C:\Users\14044\OneDrive - Conscious Care Connections\Report Quality\Database\HCP Goal Track.xlsx'
template_path = r'C:\Users\14044\OneDrive - Conscious Care Connections\Report Quality\Templates\Healthcare Plan Template.docx'
output_path = r'C:\Users\14044\OneDrive - Conscious Care Connections\Desktop\HCP Train 2024\Healthcare Plan'
treatment_db_path = r"C:\Users\14044\OneDrive - Conscious Care Connections\Desktop\HCP Train 2024\Healthcare Plan\Treatment Database.xlsx"

# Load Excel sheets
hcp_data = pd.ExcelFile(database_path)
sheet1_df = hcp_data.parse('Sheet1')  # Primary sheet for individual data
sheet2_df = hcp_data.parse('Sheet2')  # Sheet2 for general plan goals and interventions

# Clean column names: remove non-alphanumeric characters and extra spaces
def clean_column_names(df):
    df.columns = df.columns.str.strip()  # Remove leading/trailing spaces
    df.columns = [re.sub(r'\W+', '', col) for col in df.columns]  # Remove non-alphanumeric characters
    return df

# Clean column names for both Sheet1 and Sheet2
sheet1_df = clean_column_names(sheet1_df)
sheet2_df = clean_column_names(sheet2_df)

# Function to create folders and avoid duplicates
def create_folder(folder_name):
    folder_name = str(folder_name)  # Ensure folder name is a string
    folder_path = os.path.join(output_path, folder_name)
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Created folder: {folder_path}")
    return folder_path

# Function to remove individual placeholder if no data
def remove_placeholder(doc, placeholder):
    # Remove from paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, '')
    # Remove from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, '')
    # Remove from headers and footers
    for section in doc.sections:
        # Headers
        header = section.header
        for paragraph in header.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, '')
        # Footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, '')

# Function to replace placeholders
def replace_placeholders(doc, placeholder, replacement):
    if replacement == '' or pd.isna(replacement):
        remove_placeholder(doc, placeholder)
    else:
        replacement = str(replacement)
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, replacement)
        # Replace in headers and footers
        for section in doc.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)

# Function to format the date
def format_date(date_value):
    if pd.notna(date_value):
        return pd.to_datetime(date_value).strftime('%m/%d/%Y')
    return ''

# Function to format Medicaid numbers
def format_medicaid(medicaid_value):
    if pd.notna(medicaid_value):
        return str(int(medicaid_value))
    return ''

# Function to format CID
def format_cid(cid_value):
    if pd.notna(cid_value):
        return str(int(cid_value))
    return ''

# Function to process MedRecSTART data from the treatment database and insert into the document
def process_medrecstart(doc, file_name_search, treatment_sheet):
    med_rec_starts = []
    
    # Search for MedRecSTART in the treatment sheet
    for treatment_row in range(2, treatment_sheet.max_row + 1):
        treatment_file_name_search = treatment_sheet.cell(treatment_row, 1).value  # Column A (File Name Search)
        med_rec_start = treatment_sheet.cell(treatment_row, 6).value  # Column F (MedRecSTART)
        
        if treatment_file_name_search == file_name_search:
            if med_rec_start:
                med_rec_starts.append(med_rec_start)

    # Insert MedRecSTART into the document's table
    placeholder_found = False
    for table in doc.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if cell.text.strip() == "Treatments and Interventions":
                    placeholder_found = True
                    target_cell = row.cells[idx + 1] if idx + 1 < len(row.cells) else None
                    if target_cell:
                        target_cell.text = ""  # Clear the cell's content
                        for med_rec_start in med_rec_starts:
                            target_cell.add_paragraph(f'• {med_rec_start}')
                    break
        if placeholder_found:
            break
    if not placeholder_found:
        print("Header 'Treatments and Interventions' not found in the document.")

# Load treatment workbook
treatment_wb = openpyxl.load_workbook(treatment_db_path)
treatment_sheet = treatment_wb.active

# Process each entry in 'Sheet1' for every individual, applying all Healthcare Plans
hcp_names = [
    'Preventative and Routine Healthcare Maintenance Healthcare Plan',
    'Reproductive System Management Healthcare Plan',
    'Musculoskeletal Management and Falls Risk Healthcare Plan',
    'Skin Integumentary Management Healthcare Plan',
    'Neurological Management Health Care Plan',
    'Bowel and Bladder Management Healthcare Plan',
    'Neurological Management Health Care Plan'
]

for index, row in sheet1_df.iterrows():
    index_name = row['IndexName']
    if pd.isna(index_name):
        print(f"Skipping row {index} due to missing Index Name")
        continue

    index_name = str(index_name)
    for hcp_name in hcp_names:
        print(f"Processing {hcp_name} for Index Name: {index_name}")

        # Load the healthcare plan template
        doc = Document(template_path)

        # Replace placeholders in the header, footer, and main body using the correct column names from Sheet1
        replace_placeholders(doc, '<<CID>>', format_cid(row['CID']))
        replace_placeholders(doc, '<<Index Name>>', index_name)
        replace_placeholders(doc, '<<DOB>>', format_date(row['DOB']))
        replace_placeholders(doc, '<<Gender>>', row['Gender'])
        replace_placeholders(doc, '<<Race>>', row['Race'])
        replace_placeholders(doc, '<<PrimaryDx >>', row['PrimaryDx'])
        replace_placeholders(doc, '<<AdmitCode>>', row['AdmitCode'])
        replace_placeholders(doc, '<<Allergy>>', row['Allergy'])
        replace_placeholders(doc, '<<Medicaid>>', format_medicaid(row['Medicaid']))
        replace_placeholders(doc, '<<Res Pro>>', row['ResPro'])
        replace_placeholders(doc, '<<HCP Name>>', hcp_name)
        replace_placeholders(doc, '<<SVC Admit Criteria>>', row['SVCAdmitCriteria'])


        # Replace placeholders in the body with data from Sheet1 for Risk factors
        replace_placeholders(doc, '<<Risk-HIGH ALERT>>', row['RiskHIGHALERT'])
        replace_placeholders(doc, '<<Psych Med Risk>>', row['PsychMedRisk'])
        replace_placeholders(doc, '<<Cardiac Med Risk>>', row['CardiacMedRisk'])
        replace_placeholders(doc, '<<Neuro Risk>>', row['NeuroRisk'])
        replace_placeholders(doc, '<<Aspiration Risk>>', row['AspirationRisk'])
        
        # Replace placeholders in the body with data from Sheet2 (goals and interventions)
        hcp_row = sheet2_df[sheet2_df['HCPName'] == hcp_name]  # Match the correct HCP name from Sheet2
        
        if not hcp_row.empty:
            hcp_data = hcp_row.iloc[0]  # Get the first matching row for the HCP Name
            replace_placeholders(doc, '<<HCPGoal>>', hcp_data['HCPGoal'])
            replace_placeholders(doc, '<<HCPGoal2>>', hcp_data['HCPGoal2'])
            replace_placeholders(doc, '<<HMA1>>', hcp_data['HMA1'])
            replace_placeholders(doc, '<<HMA2>>', hcp_data['HMA2'])
            replace_placeholders(doc, '<<HMA3>>', hcp_data['HMA3'])
            replace_placeholders(doc, '<<HMA4>>', hcp_data['HMA4'])
            replace_placeholders(doc, '<<HMA5>>', hcp_data['HMA5'])
            replace_placeholders(doc, '<<HTrack1>>', hcp_data['HTrack1'])
            replace_placeholders(doc, '<<HTrack2>>', hcp_data['HTrack2'])
        else:
            print(f"No matching HCP Name found in Sheet2 for {hcp_name}. Removing placeholders.")
            placeholders_to_remove = ['<<HCPGoal>>', '<<HCPGoal2>>', '<<HMA1>>', '<<HMA2>>', '<<HMA3>>', '<<HMA4>>', '<<HMA5>>', '<<HTrack1>>', '<<HTrack2>>']
            for placeholder in placeholders_to_remove:
                remove_placeholder(doc, placeholder)
        # Replace remaining placeholders if needed (e.g., <<SVC Admit Criteria>>)
        replace_placeholders(doc, '<<SVC Admit Criteria>>', row['SVCAdmitCriteria'])
       
        
        # Save the document for this HCP with unique identifier, index name, and HCP name
        folder_path = create_folder(index_name)
        unique_identifier = str(row['UniqueIdentifier']) if pd.notna(row['UniqueIdentifier']) else 'Unknown'
        output_filename = f"{unique_identifier}_{index_name}_{hcp_name}.docx"
        doc.save(os.path.join(folder_path, output_filename))


        # Process MedRecSTART based on the file name search
        file_name_search = f"{row['UniqueIdentifier']}_{index_name}_{hcp_name}.docx"
        process_medrecstart(doc, file_name_search, treatment_sheet)

        # Save the document after all processing
        folder_path = create_folder(index_name)
        output_file_path = os.path.join(folder_path, file_name_search)
        doc.save(output_file_path)
        print(f"Document saved: {output_file_path}")
	
# Call garbage collector
        gc.collect()

# Close the treatment workbook
treatment_wb.close()
print("Healthcare plans updated successfully.")
